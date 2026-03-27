import asyncio
import io
import os
import sqlite3
import urllib.error
from uuid import uuid4
from typing import Annotated, Optional, TypedDict

import fitz  # PyMuPDF
import openpyxl
from docx import Document as DocxDocument
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langgraph.checkpoint.sqlite import SqliteSaver
from langgraph.graph import END, StateGraph
from langgraph.graph.message import add_messages
from pydantic import BaseModel

try:
    from api_errors import error_detail
    from artifacts import (
        ProjectArtifacts,
        delivery_items_to_json,
        export_project_json,
        export_project_markdown,
    )
    from integrations.jira import list_jira_projects
    from integrations.github import list_github_repos
    from integrations.registry_map import DELIVERY_INTEGRATIONS
    from model_adapters import MODEL_ADAPTERS, get_supported_model_choices, invoke_model
    from prompts import (
        build_arch_chat_prompt,
        build_architecture_refine_prompt,
        build_architect_prompt,
        build_prd_refine_prompt,
        build_sa_prompt,
        build_stories_chat_prompt,
        build_user_stories_refine_prompt,
        build_user_stories_prompt,
    )
    from workflow import parse_delivery_items, project_artifacts_from_state
except ModuleNotFoundError:
    from backend.api_errors import error_detail
    from backend.artifacts import (
        ProjectArtifacts,
        delivery_items_to_json,
        export_project_json,
        export_project_markdown,
    )
    from backend.integrations.jira import list_jira_projects
    from backend.integrations.github import list_github_repos
    from backend.integrations.registry_map import DELIVERY_INTEGRATIONS
    from backend.model_adapters import MODEL_ADAPTERS, get_supported_model_choices, invoke_model
    from backend.prompts import (
        build_arch_chat_prompt,
        build_architecture_refine_prompt,
        build_architect_prompt,
        build_prd_refine_prompt,
        build_sa_prompt,
        build_stories_chat_prompt,
        build_user_stories_refine_prompt,
        build_user_stories_prompt,
    )
    from backend.workflow import parse_delivery_items, project_artifacts_from_state

CORS_ALLOW_ORIGINS = [
    origin.strip()
    for origin in os.getenv("CORS_ALLOW_ORIGINS", "*").split(",")
    if origin.strip()
]


# ---------------------------------------------------------------------------
# SQLite persistence
# ---------------------------------------------------------------------------

conn = sqlite3.connect("ai_factory.db", check_same_thread=False)
memory = SqliteSaver(conn)
memory.setup()


# ---------------------------------------------------------------------------
# LangGraph State
# ---------------------------------------------------------------------------

class SAState(TypedDict):
    messages: Annotated[list, add_messages]
    prd_draft: str
    is_ready_for_architecture: bool
    model_choice: str
    architecture_draft: str
    user_stories_draft: str


# ---------------------------------------------------------------------------
# Stage chat — SQLite-backed message store
# ---------------------------------------------------------------------------

CONTENT_START_MARKER = "[CONTENT_START]"
CONTENT_END_MARKER   = "[CONTENT_END]"

conn.execute("""
    CREATE TABLE IF NOT EXISTS projects (
        thread_id  TEXT PRIMARY KEY,
        name       TEXT NOT NULL,
        created_at REAL NOT NULL DEFAULT (strftime('%s','now'))
    )
""")
conn.execute(
    "CREATE INDEX IF NOT EXISTS idx_projects_created_at ON projects (created_at, thread_id)"
)
conn.execute("""
    CREATE TABLE IF NOT EXISTS stage_messages (
        id         INTEGER PRIMARY KEY AUTOINCREMENT,
        thread_id  TEXT    NOT NULL,
        stage      TEXT    NOT NULL,
        role       TEXT    NOT NULL,
        content    TEXT    NOT NULL,
        created_at REAL    NOT NULL DEFAULT (strftime('%s','now'))
    )
""")
conn.execute(
    "CREATE INDEX IF NOT EXISTS idx_stage_messages ON stage_messages (thread_id, stage, id)"
)
conn.execute("""
    CREATE TABLE IF NOT EXISTS stage_status (
        thread_id  TEXT NOT NULL,
        stage      TEXT NOT NULL,
        status     TEXT NOT NULL DEFAULT 'draft',
        updated_at REAL NOT NULL DEFAULT (strftime('%s','now')),
        PRIMARY KEY (thread_id, stage)
    )
""")
conn.execute("""
    CREATE TABLE IF NOT EXISTS stage_events (
        id         INTEGER PRIMARY KEY AUTOINCREMENT,
        thread_id  TEXT    NOT NULL,
        stage      TEXT    NOT NULL,
        event_type TEXT    NOT NULL,
        detail     TEXT,
        created_at REAL    NOT NULL DEFAULT (strftime('%s','now'))
    )
""")
conn.execute(
    "CREATE INDEX IF NOT EXISTS idx_stage_events ON stage_events (thread_id, stage, id)"
)
conn.execute("""
    CREATE TABLE IF NOT EXISTS stage_comments (
        id          INTEGER PRIMARY KEY AUTOINCREMENT,
        thread_id   TEXT    NOT NULL,
        stage       TEXT    NOT NULL,
        body        TEXT    NOT NULL,
        status      TEXT    NOT NULL DEFAULT 'open',
        created_at  REAL    NOT NULL DEFAULT (strftime('%s','now')),
        resolved_at REAL
    )
""")
conn.execute(
    "CREATE INDEX IF NOT EXISTS idx_stage_comments ON stage_comments (thread_id, stage, id)"
)
conn.execute("""
    CREATE TABLE IF NOT EXISTS stage_revisions (
        id               INTEGER PRIMARY KEY AUTOINCREMENT,
        thread_id        TEXT    NOT NULL,
        stage            TEXT    NOT NULL,
        source           TEXT    NOT NULL,
        summary          TEXT    NOT NULL DEFAULT '',
        instruction      TEXT    NOT NULL DEFAULT '',
        reviewed         INTEGER NOT NULL DEFAULT 0,
        downstream_reset TEXT    NOT NULL DEFAULT '',
        content_length   INTEGER NOT NULL DEFAULT 0,
        created_at       REAL    NOT NULL DEFAULT (strftime('%s','now'))
    )
""")
conn.execute(
    "CREATE INDEX IF NOT EXISTS idx_stage_revisions ON stage_revisions (thread_id, stage, id)"
)
conn.commit()

STAGE_KEYS = ("prd", "architecture", "stories")
STAGE_DEPENDENCIES: dict[str, list[str]] = {
    "prd": [],
    "architecture": ["prd"],
    "stories": ["architecture"],
}
STAGE_DOWNSTREAM: dict[str, list[str]] = {
    "prd": ["architecture", "stories"],
    "architecture": ["stories"],
    "stories": [],
}


def _get_stage_status(thread_id: str, stage: str) -> str:
    row = conn.execute(
        "SELECT status FROM stage_status WHERE thread_id=? AND stage=?",
        (thread_id, stage),
    ).fetchone()
    return row[0] if row else "draft"


def _set_stage_status(thread_id: str, stage: str, status: str) -> None:
    conn.execute(
        """INSERT INTO stage_status (thread_id, stage, status, updated_at)
           VALUES (?, ?, ?, strftime('%s','now'))
           ON CONFLICT(thread_id, stage) DO UPDATE SET
               status=excluded.status, updated_at=excluded.updated_at""",
        (thread_id, stage, status),
    )
    conn.commit()


def _reset_stage_status(thread_id: str, stage: str) -> None:
    """Reset to draft when content is regenerated."""
    _set_stage_status(thread_id, stage, "draft")


def _record_stage_event(thread_id: str, stage: str, event_type: str, detail: str = "") -> None:
    conn.execute(
        "INSERT INTO stage_events (thread_id, stage, event_type, detail) VALUES (?, ?, ?, ?)",
        (thread_id, stage, event_type, detail or ""),
    )
    conn.commit()


def _load_stage_events(thread_id: str, stage: str) -> list[dict]:
    rows = conn.execute(
        "SELECT event_type, detail, created_at FROM stage_events WHERE thread_id=? AND stage=? ORDER BY id DESC LIMIT 50",
        (thread_id, stage),
    ).fetchall()
    return [{"event_type": row[0], "detail": row[1], "created_at": row[2]} for row in rows]


def _record_stage_revision(
    thread_id: str,
    stage: str,
    source: str,
    *,
    summary: str = "",
    instruction: str = "",
    reviewed: bool = False,
    downstream_reset: Optional[list[str]] = None,
    content_length: int = 0,
) -> None:
    conn.execute(
        """
        INSERT INTO stage_revisions (
            thread_id, stage, source, summary, instruction, reviewed, downstream_reset, content_length
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            thread_id,
            stage,
            source,
            summary.strip(),
            instruction.strip(),
            1 if reviewed else 0,
            ",".join(downstream_reset or []),
            content_length,
        ),
    )
    conn.commit()


def _load_stage_revisions(thread_id: str, stage: str, limit: int = 20) -> list[dict]:
    rows = conn.execute(
        """
        SELECT id, source, summary, instruction, reviewed, downstream_reset, content_length, created_at
        FROM stage_revisions
        WHERE thread_id=? AND stage=?
        ORDER BY id DESC
        LIMIT ?
        """,
        (thread_id, stage, limit),
    ).fetchall()
    return [
        {
            "id": row[0],
            "source": row[1],
            "summary": row[2],
            "instruction": row[3],
            "reviewed": bool(row[4]),
            "downstream_reset": [item for item in row[5].split(",") if item],
            "content_length": row[6],
            "created_at": row[7],
        }
        for row in rows
    ]


def _get_latest_stage_revision(thread_id: str, stage: str) -> Optional[dict]:
    revisions = _load_stage_revisions(thread_id, stage, limit=1)
    return revisions[0] if revisions else None


def _get_latest_stage_timestamp(thread_id: str, stage: str) -> Optional[float]:
    revision_row = conn.execute(
        "SELECT created_at FROM stage_revisions WHERE thread_id=? AND stage=? ORDER BY id DESC LIMIT 1",
        (thread_id, stage),
    ).fetchone()
    if revision_row:
        return float(revision_row[0])
    event_row = conn.execute(
        "SELECT created_at FROM stage_events WHERE thread_id=? AND stage=? ORDER BY id DESC LIMIT 1",
        (thread_id, stage),
    ).fetchone()
    if event_row:
        return float(event_row[0])
    return None


def _load_stage_comments(thread_id: str, stage: str) -> list[dict]:
    rows = conn.execute(
        """
        SELECT id, body, status, created_at, resolved_at
        FROM stage_comments
        WHERE thread_id=? AND stage=?
        ORDER BY
            CASE status WHEN 'open' THEN 0 ELSE 1 END,
            id DESC
        """,
        (thread_id, stage),
    ).fetchall()
    return [
        {
            "id": row[0],
            "body": row[1],
            "status": row[2],
            "created_at": row[3],
            "resolved_at": row[4],
        }
        for row in rows
    ]


def _create_stage_comment(thread_id: str, stage: str, body: str) -> dict:
    cursor = conn.execute(
        "INSERT INTO stage_comments (thread_id, stage, body) VALUES (?, ?, ?)",
        (thread_id, stage, body.strip()),
    )
    conn.commit()
    comment_id = cursor.lastrowid
    row = conn.execute(
        """
        SELECT id, body, status, created_at, resolved_at
        FROM stage_comments
        WHERE id=?
        """,
        (comment_id,),
    ).fetchone()
    return {
        "id": row[0],
        "body": row[1],
        "status": row[2],
        "created_at": row[3],
        "resolved_at": row[4],
    }


def _update_stage_comment_status(comment_id: int, status: str) -> Optional[dict]:
    resolved_at_sql = "strftime('%s','now')" if status == "resolved" else "NULL"
    conn.execute(
        f"UPDATE stage_comments SET status=?, resolved_at={resolved_at_sql} WHERE id=?",
        (status, comment_id),
    )
    conn.commit()
    row = conn.execute(
        """
        SELECT id, thread_id, stage, body, status, created_at, resolved_at
        FROM stage_comments
        WHERE id=?
        """,
        (comment_id,),
    ).fetchone()
    if not row:
        return None
    return {
        "id": row[0],
        "thread_id": row[1],
        "stage": row[2],
        "body": row[3],
        "status": row[4],
        "created_at": row[5],
        "resolved_at": row[6],
    }


def _load_stage_messages(thread_id: str, stage: str) -> list[dict]:
    rows = conn.execute(
        "SELECT role, content FROM stage_messages WHERE thread_id=? AND stage=? ORDER BY id",
        (thread_id, stage),
    ).fetchall()
    return [{"role": row[0], "content": row[1]} for row in rows]


def _append_stage_message(thread_id: str, stage: str, role: str, content: str) -> None:
    conn.execute(
        "INSERT INTO stage_messages (thread_id, stage, role, content) VALUES (?, ?, ?, ?)",
        (thread_id, stage, role, content),
    )
    conn.commit()


def _delete_stage_messages(thread_id: str) -> None:
    conn.execute("DELETE FROM stage_messages WHERE thread_id=?", (thread_id,))
    conn.execute("DELETE FROM stage_status WHERE thread_id=?", (thread_id,))
    conn.execute("DELETE FROM stage_events WHERE thread_id=?", (thread_id,))
    conn.execute("DELETE FROM stage_comments WHERE thread_id=?", (thread_id,))
    conn.execute("DELETE FROM stage_revisions WHERE thread_id=?", (thread_id,))
    conn.commit()


def _list_projects() -> list[dict[str, str]]:
    rows = conn.execute(
        "SELECT thread_id, name FROM projects ORDER BY created_at ASC, thread_id ASC"
    ).fetchall()
    return [{"id": row[0], "name": row[1]} for row in rows]


def _upsert_project(thread_id: str, name: str) -> None:
    conn.execute(
        """INSERT INTO projects (thread_id, name, created_at)
           VALUES (?, ?, strftime('%s','now'))
           ON CONFLICT(thread_id) DO UPDATE SET name=excluded.name""",
        (thread_id, name),
    )
    conn.commit()


def _delete_project_record(thread_id: str) -> None:
    conn.execute("DELETE FROM projects WHERE thread_id=?", (thread_id,))
    conn.commit()


def _delete_thread_state(thread_id: str) -> None:
    for table in ("checkpoints", "checkpoint_writes"):
        try:
            conn.execute(f"DELETE FROM {table} WHERE thread_id = ?", (thread_id,))
            conn.commit()
        except sqlite3.Error:
            pass
    _delete_stage_messages(thread_id)
    _delete_project_record(thread_id)


def _normalize_change_source(change_source: str) -> str:
    normalized = change_source.strip().lower() if change_source else "manual_edit"
    if normalized not in {"manual_edit", "ai_revision"}:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_input",
                "change_source must be one of: manual_edit, ai_revision.",
            ),
        )
    return normalized


def _event_type_for_change_source(change_source: str) -> str:
    return "ai_revised" if change_source == "ai_revision" else "manually_edited"


def _default_revision_summary(stage: str, change_source: str, change_context: str) -> str:
    context = (change_context or "").strip().lower()
    if context == "stage_chat":
        return f"{stage.title()} updated from stage discussion"
    if context == "ai_refine":
        return f"{stage.title()} updated from AI revision"
    if context == "manual_edit":
        return f"{stage.title()} updated manually"
    return f"{stage.title()} updated by {'AI' if change_source == 'ai_revision' else 'manual edit'}"


def _stage_has_content(artifacts: ProjectArtifacts, stage: str) -> bool:
    if stage == "prd":
        return bool(artifacts.prd.strip())
    if stage == "architecture":
        return bool(artifacts.architecture.strip())
    if stage == "stories":
        return bool(artifacts.user_stories.strip())
    return False


def _extract_stage_content(response_text: str) -> "tuple[str, Optional[str]]":
    """
    Split agent response into (conversational_part, updated_artifact | None).

    If the agent wrapped updated content in [CONTENT_START]…[CONTENT_END],
    the text before the marker becomes the chat message and the content inside
    becomes the updated artifact.  Otherwise the full text is the chat message.
    """
    start = response_text.find(CONTENT_START_MARKER)
    end   = response_text.find(CONTENT_END_MARKER)
    if start != -1 and end != -1 and end > start:
        conversation_part = response_text[:start].strip()
        content_part      = response_text[start + len(CONTENT_START_MARKER):end].strip()
        return conversation_part or "I've updated the content as requested.", content_part
    return response_text, None


# ---------------------------------------------------------------------------
# SA Interaction Node
# ---------------------------------------------------------------------------

def sa_interaction_node(state: SAState) -> SAState:
    """
    System Analyst agent node.

    Builds a full conversation prompt (system + history) and dispatches it
    to the model selected in state["model_choice"].

    If a finalized PRD already exists (is_ready_for_architecture is True),
    the node runs in Amendment Mode: it injects the current PRD into the prompt
    so the agent knows it must produce an updated, complete PRD that merges
    original + new requirements.
    """
    model_choice = state.get("model_choice", "ollama")
    existing_prd = state.get("prd_draft", "")
    already_ready = state.get("is_ready_for_architecture", False)

    # Build a single string prompt that concatenates the system instructions
    # and the conversation history — works for all adapters.
    history_lines: list[str] = []
    for msg in state["messages"]:
        if isinstance(msg, HumanMessage):
            history_lines.append(f"User: {msg.content}")
        elif isinstance(msg, AIMessage):
            history_lines.append(f"SA Agent: {msg.content}")
        elif isinstance(msg, SystemMessage):
            pass  # handled separately below

    conversation_text = "\n".join(history_lines)

    full_prompt = build_sa_prompt(
        model_choice=model_choice,
        conversation_text=conversation_text,
        existing_prd=existing_prd,
        already_ready=already_ready,
    )

    try:
        response_text = invoke_model(model_choice, full_prompt)
    except (RuntimeError, ValueError) as exc:
        response_text = (
            f"[ModelAdapter Error] Unable to get a response from '{model_choice}': "
            f"{exc}\n\n"
            "Please check that the selected AI backend is installed and running, "
            "then try again."
        )

    prd_draft = state.get("prd_draft", "")
    is_ready = state.get("is_ready_for_architecture", False)

    if "[PRD_READY]" in response_text:
        is_ready = True
        prd_draft = response_text.replace("[PRD_READY]", "").rstrip()

    return {
        "messages": [AIMessage(content=response_text)],
        "prd_draft": prd_draft,
        "is_ready_for_architecture": is_ready,
        "model_choice": model_choice,
        "architecture_draft": state.get("architecture_draft", ""),
        "user_stories_draft": state.get("user_stories_draft", ""),
    }


# ---------------------------------------------------------------------------
# LangGraph setup
# ---------------------------------------------------------------------------

graph_builder = StateGraph(SAState)
graph_builder.add_node("sa_agent", sa_interaction_node)
graph_builder.set_entry_point("sa_agent")
graph_builder.add_edge("sa_agent", END)

graph = graph_builder.compile(checkpointer=memory)


# ---------------------------------------------------------------------------
# FastAPI
# ---------------------------------------------------------------------------

app = FastAPI(
    title="Interactive AI Software Factory",
    description="HITL AI Software Factory — Dynamic Model Adapter (Ollama / Gemini CLI / Claude CLI / Codex CLI) with System Analyst Agent",
    version="2.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOW_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

VALID_MODEL_CHOICES = set(get_supported_model_choices())
VALID_DELIVERY_TARGETS = set(DELIVERY_INTEGRATIONS.keys())


def validate_model_choice(model_choice: str) -> str:
    normalized = model_choice.strip().lower()
    if normalized not in VALID_MODEL_CHOICES:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_model_choice",
                f"Invalid model_choice '{normalized}'. "
                f"Valid options: {sorted(VALID_MODEL_CHOICES)}",
            ),
        )
    return normalized


def validate_delivery_target(target: str) -> str:
    normalized = target.strip().lower()
    if normalized not in VALID_DELIVERY_TARGETS:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_delivery_target",
                f"Invalid delivery target '{normalized}'. "
                f"Valid options: {sorted(VALID_DELIVERY_TARGETS)}",
            ),
        )
    return normalized


def get_thread_values(thread_id: str) -> dict:
    config = {"configurable": {"thread_id": thread_id}}
    try:
        state_snapshot = graph.get_state(config)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to retrieve thread state: {exc}"),
        ) from exc

    if state_snapshot is None or not state_snapshot.values:
        return {}
    return state_snapshot.values


def get_project_artifacts(thread_id: str) -> ProjectArtifacts:
    return project_artifacts_from_state(thread_id, get_thread_values(thread_id))


def require_user_stories(thread_id: str) -> str:
    artifacts = get_project_artifacts(thread_id)
    if not artifacts.user_stories:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "missing_user_stories",
                "User stories must be generated before using delivery integrations.",
            ),
        )
    return artifacts.user_stories


def require_prd(thread_id: str) -> str:
    artifacts = get_project_artifacts(thread_id)
    if not artifacts.prd:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "missing_prd",
                "PRD must exist before this stage can be updated.",
            ),
        )
    return artifacts.prd


def require_architecture(thread_id: str) -> str:
    artifacts = get_project_artifacts(thread_id)
    if not artifacts.architecture:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "missing_architecture",
                "Architecture must exist before this stage can be updated.",
            ),
        )
    return artifacts.architecture


def parse_delivery_items_or_raise(thread_id: str, model_choice: str):
    user_stories_draft = require_user_stories(thread_id)
    try:
        return parse_delivery_items(user_stories_draft, model_choice)
    except ValueError as exc:
        raise HTTPException(
            status_code=422,
            detail=error_detail(
                "story_parse_error",
                f"Could not parse model output as JSON: {exc}",
            ),
        ) from exc
    except (RuntimeError, FileNotFoundError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"Model failed to parse user stories: {exc}"),
        ) from exc


def build_delivery_config(
    target: str,
    *,
    require_credentials: bool = True,
    jira_domain: str = "",
    jira_email: str = "",
    jira_token: str = "",
    jira_project_key: str = "",
    github_owner: str = "",
    github_repo: str = "",
    github_token: str = "",
) -> dict[str, str]:
    if target == "jira":
        if not jira_project_key.strip():
            raise HTTPException(
                status_code=400,
                detail=error_detail(
                    "invalid_integration_config",
                    "Jira delivery requires a project key.",
                ),
            )
        if require_credentials and not all([jira_domain.strip(), jira_email.strip(), jira_token.strip()]):
            raise HTTPException(
                status_code=400,
                detail=error_detail(
                    "invalid_integration_config",
                    "Jira publish requires domain, email, and API token.",
                ),
            )
        return {
            "domain": jira_domain.strip(),
            "email": jira_email.strip(),
            "token": jira_token.strip(),
            "project_key": jira_project_key.strip(),
        }

    if target == "github":
        if not all([github_owner.strip(), github_repo.strip()]):
            raise HTTPException(
                status_code=400,
                detail=error_detail(
                    "invalid_integration_config",
                    "GitHub delivery requires owner and repo.",
                ),
            )
        if require_credentials and not github_token.strip():
            raise HTTPException(
                status_code=400,
                detail=error_detail(
                    "invalid_integration_config",
                    "GitHub publish requires a personal access token.",
                ),
            )
        return {
            "owner": github_owner.strip(),
            "repo": github_repo.strip(),
            "token": github_token.strip(),
        }

    raise HTTPException(
        status_code=400,
        detail=error_detail("invalid_delivery_target", f"Unsupported target '{target}'."),
    )


class ChatRequest(BaseModel):
    thread_id: str
    user_input: str
    model_choice: str = "ollama"


class ChatResponse(BaseModel):
    ai_response: str
    current_prd: str
    is_ready: bool
    model_used: str


@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    if not request.user_input.strip():
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "user_input cannot be empty."),
        )
    if not request.thread_id.strip():
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "thread_id cannot be empty."),
        )

    model_choice = request.model_choice.strip().lower()
    if model_choice not in VALID_MODEL_CHOICES:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_model_choice",
                f"Invalid model_choice '{model_choice}'. "
                f"Valid options: {sorted(VALID_MODEL_CHOICES)}",
            ),
        )

    config = {"configurable": {"thread_id": request.thread_id}}

    # Only supply fields that are new for this turn.
    # Omitting prd_draft, is_ready_for_architecture, architecture_draft, and
    # user_stories_draft preserves their values from the existing checkpoint
    # so the SA Agent can read — and build upon — previously accumulated state.
    input_state = {
        "messages": [HumanMessage(content=request.user_input)],
        "model_choice": model_choice,
    }

    try:
        result = graph.invoke(input_state, config=config)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("workflow_error", f"LangGraph invocation failed: {exc}"),
        )

    ai_messages = [m for m in result["messages"] if isinstance(m, AIMessage)]
    if not ai_messages:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", "No AI response was generated."),
        )

    latest_response = ai_messages[-1].content

    return ChatResponse(
        ai_response=latest_response,
        current_prd=result.get("prd_draft", ""),
        is_ready=result.get("is_ready_for_architecture", False),
        model_used=model_choice,
    )


class ThreadMessage(BaseModel):
    role: str
    content: str


class ThreadStateResponse(BaseModel):
    messages: list[ThreadMessage]
    current_prd: str
    is_ready: bool
    architecture_draft: str
    user_stories_draft: str


class ProjectSummary(BaseModel):
    id: str
    name: str


class ProjectsListResponse(BaseModel):
    projects: list[ProjectSummary]


class CreateProjectRequest(BaseModel):
    name: str
    thread_id: Optional[str] = None


class CreateProjectResponse(BaseModel):
    project: ProjectSummary


@app.get("/api/projects", response_model=ProjectsListResponse)
async def list_projects():
    return ProjectsListResponse(
        projects=[ProjectSummary(**project) for project in _list_projects()]
    )


@app.post("/api/projects", response_model=CreateProjectResponse)
async def create_project(request: CreateProjectRequest):
    name = request.name.strip()
    if not name:
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "Project name cannot be empty."),
        )

    thread_id = (request.thread_id or "").strip() or uuid4().hex
    try:
        _upsert_project(thread_id, name)
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to persist project metadata: {exc}"),
        ) from exc

    return CreateProjectResponse(project=ProjectSummary(id=thread_id, name=name))


@app.get("/api/chat/{thread_id}", response_model=ThreadStateResponse)
async def get_thread_state(thread_id: str):
    """Return the persisted conversation state for a given thread_id."""
    config = {"configurable": {"thread_id": thread_id}}
    try:
        state_snapshot = graph.get_state(config)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to retrieve thread state: {exc}"),
        )

    # No checkpoint exists yet for this thread
    if state_snapshot is None or not state_snapshot.values:
        return ThreadStateResponse(
            messages=[], current_prd="", is_ready=False, architecture_draft="", user_stories_draft=""
        )

    values = state_snapshot.values
    raw_messages = values.get("messages", [])

    thread_messages: list[ThreadMessage] = []
    for msg in raw_messages:
        if isinstance(msg, HumanMessage):
            thread_messages.append(ThreadMessage(role="user", content=msg.content))
        elif isinstance(msg, AIMessage):
            thread_messages.append(ThreadMessage(role="assistant", content=msg.content))

    return ThreadStateResponse(
        messages=thread_messages,
        current_prd=values.get("prd_draft", ""),
        is_ready=values.get("is_ready_for_architecture", False),
        architecture_draft=values.get("architecture_draft", ""),
        user_stories_draft=values.get("user_stories_draft", ""),
    )


class UpdatePrdRequest(BaseModel):
    content: str
    change_source: str = "manual_edit"
    reviewed: bool = False
    instruction: str = ""
    change_context: str = ""


class UpdatePrdResponse(BaseModel):
    success: bool
    prd_draft: str
    is_ready: bool


@app.put("/api/prd/{thread_id}", response_model=UpdatePrdResponse)
async def update_prd(thread_id: str, request: UpdatePrdRequest):
    content = request.content.strip()
    if not content:
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "PRD content cannot be empty."),
        )
    change_source = _normalize_change_source(request.change_source)

    config = {"configurable": {"thread_id": thread_id}}
    try:
        graph.update_state(
            config,
            {
                "prd_draft": content,
                "is_ready_for_architecture": True,
                "architecture_draft": "",
                "user_stories_draft": "",
            },
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail(
                "state_error",
                f"Failed to persist PRD for thread '{thread_id}': {exc}",
            ),
        )
    _reset_stage_status(thread_id, "prd")
    _reset_stage_status(thread_id, "architecture")
    _reset_stage_status(thread_id, "stories")
    _record_stage_event(thread_id, "prd", _event_type_for_change_source(change_source))
    _record_stage_revision(
        thread_id,
        "prd",
        change_source,
        summary=_default_revision_summary("prd", change_source, request.change_context),
        instruction=request.instruction,
        reviewed=request.reviewed,
        downstream_reset=STAGE_DOWNSTREAM["prd"],
        content_length=len(content),
    )
    return UpdatePrdResponse(success=True, prd_draft=content, is_ready=True)


class RefinePrdRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    instruction: str
    preview_only: bool = False


class RefinePrdResponse(BaseModel):
    prd_draft: str
    is_ready: bool


@app.post("/api/refine_prd", response_model=RefinePrdResponse)
async def refine_prd(request: RefinePrdRequest):
    if not request.instruction.strip():
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "instruction cannot be empty."),
        )

    model_choice = validate_model_choice(request.model_choice)
    existing_prd = require_prd(request.thread_id)
    prompt = build_prd_refine_prompt(
        model_choice=model_choice,
        prd_draft=existing_prd,
        instruction=request.instruction.strip(),
    )

    try:
        result = await asyncio.to_thread(invoke_model, model_choice, prompt)
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"PRD refine failed: {exc}"),
        )

    updated_prd = result.replace("[PRD_READY]", "").rstrip()
    if request.preview_only:
        return RefinePrdResponse(prd_draft=updated_prd, is_ready=True)

    config = {"configurable": {"thread_id": request.thread_id}}
    try:
        graph.update_state(
            config,
            {
                "prd_draft": updated_prd,
                "is_ready_for_architecture": True,
                "architecture_draft": "",
                "user_stories_draft": "",
            },
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to persist refined PRD: {exc}"),
        )

    _reset_stage_status(request.thread_id, "prd")
    _reset_stage_status(request.thread_id, "architecture")
    _reset_stage_status(request.thread_id, "stories")
    _record_stage_event(request.thread_id, "prd", "ai_revised")
    _record_stage_revision(
        request.thread_id,
        "prd",
        "ai_revision",
        summary="PRD updated from AI revision",
        instruction=request.instruction,
        reviewed=not request.preview_only,
        downstream_reset=STAGE_DOWNSTREAM["prd"],
        content_length=len(updated_prd),
    )
    return RefinePrdResponse(prd_draft=updated_prd, is_ready=True)


class DeleteThreadResponse(BaseModel):
    success: bool
    thread_id: str


@app.delete("/api/chat/{thread_id}", response_model=DeleteThreadResponse)
async def delete_thread(thread_id: str):
    """Delete all persisted checkpoint data for a given thread_id."""
    try:
        _delete_thread_state(thread_id)
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to delete thread '{thread_id}': {exc}"),
        )
    return DeleteThreadResponse(success=True, thread_id=thread_id)


@app.delete("/api/projects/{thread_id}", response_model=DeleteThreadResponse)
async def delete_project(thread_id: str):
    try:
        _delete_thread_state(thread_id)
    except sqlite3.Error as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to delete project '{thread_id}': {exc}"),
        ) from exc
    return DeleteThreadResponse(success=True, thread_id=thread_id)


# ---------------------------------------------------------------------------
# Architecture generation endpoint
# ---------------------------------------------------------------------------

class GenerateArchitectureRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"


class GenerateArchitectureResponse(BaseModel):
    architecture_draft: str


@app.post("/api/generate_architecture", response_model=GenerateArchitectureResponse)
async def generate_architecture(request: GenerateArchitectureRequest):
    """
    Human-in-the-Loop gate: user approves the PRD, then this endpoint invokes
    the Architect Agent to produce a technical design with Mermaid diagrams.
    """
    model_choice = request.model_choice.strip().lower()
    if model_choice not in VALID_MODEL_CHOICES:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_model_choice",
                f"Invalid model_choice '{model_choice}'. "
                f"Valid options: {sorted(VALID_MODEL_CHOICES)}",
            ),
        )

    config = {"configurable": {"thread_id": request.thread_id}}

    # Retrieve current state
    try:
        state_snapshot = graph.get_state(config)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to retrieve thread state: {exc}"),
        )

    if state_snapshot is None or not state_snapshot.values:
        raise HTTPException(
            status_code=400,
            detail=error_detail("missing_prd", "PRD is not ready yet."),
        )

    prd_draft = state_snapshot.values.get("prd_draft", "").strip()
    if not prd_draft:
        raise HTTPException(
            status_code=400,
            detail=error_detail("missing_prd", "PRD is not ready yet."),
        )

    architect_prompt = build_architect_prompt(
        model_choice=model_choice,
        prd_draft=prd_draft,
    )

    # Invoke the model (runs in a thread to avoid blocking the event loop)
    try:
        result = await asyncio.to_thread(invoke_model, model_choice, architect_prompt)
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"Architect agent failed: {exc}"),
        )

    # Persist the architecture draft back into the graph state
    try:
        graph.update_state(config, {"architecture_draft": result})
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to persist architecture draft: {exc}"),
        )

    _reset_stage_status(request.thread_id, "architecture")
    _reset_stage_status(request.thread_id, "stories")
    _record_stage_event(request.thread_id, "architecture", "generated")
    _record_stage_revision(
        request.thread_id,
        "architecture",
        "generated",
        summary="Architecture generated from approved PRD",
        reviewed=False,
        downstream_reset=STAGE_DOWNSTREAM["architecture"],
        content_length=len(result),
    )
    return GenerateArchitectureResponse(architecture_draft=result)


# ---------------------------------------------------------------------------
# Architecture edit endpoint
# ---------------------------------------------------------------------------

class UpdateArchitectureRequest(BaseModel):
    content: str
    change_source: str = "manual_edit"
    reviewed: bool = False
    instruction: str = ""
    change_context: str = ""


class UpdateArchitectureResponse(BaseModel):
    success: bool
    architecture_draft: str


@app.put("/api/architecture/{thread_id}", response_model=UpdateArchitectureResponse)
async def update_architecture(thread_id: str, request: UpdateArchitectureRequest):
    """
    Persist a manually-edited architecture draft.
    Also clears user_stories_draft since it would be stale after an architecture edit.
    """
    change_source = _normalize_change_source(request.change_source)
    config = {"configurable": {"thread_id": thread_id}}
    try:
        graph.update_state(
            config,
            {
                "architecture_draft": request.content,
                "user_stories_draft": "",
            },
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail(
                "state_error",
                f"Failed to persist architecture for thread '{thread_id}': {exc}",
            ),
        )
    _reset_stage_status(thread_id, "architecture")
    _reset_stage_status(thread_id, "stories")
    _record_stage_event(thread_id, "architecture", _event_type_for_change_source(change_source))
    _record_stage_revision(
        thread_id,
        "architecture",
        change_source,
        summary=_default_revision_summary("architecture", change_source, request.change_context),
        instruction=request.instruction,
        reviewed=request.reviewed,
        downstream_reset=STAGE_DOWNSTREAM["architecture"],
        content_length=len(request.content),
    )
    return UpdateArchitectureResponse(success=True, architecture_draft=request.content)


class RefineArchitectureRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    instruction: str
    preview_only: bool = False


class RefineArchitectureResponse(BaseModel):
    architecture_draft: str


@app.post("/api/refine_architecture", response_model=RefineArchitectureResponse)
async def refine_architecture(request: RefineArchitectureRequest):
    if not request.instruction.strip():
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "instruction cannot be empty."),
        )

    model_choice = validate_model_choice(request.model_choice)
    prd_draft = require_prd(request.thread_id)
    architecture_draft = require_architecture(request.thread_id)
    prompt = build_architecture_refine_prompt(
        model_choice=model_choice,
        prd_draft=prd_draft,
        architecture_draft=architecture_draft,
        instruction=request.instruction.strip(),
    )

    try:
        result = await asyncio.to_thread(invoke_model, model_choice, prompt)
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"Architecture refine failed: {exc}"),
        )

    if request.preview_only:
        return RefineArchitectureResponse(architecture_draft=result)

    config = {"configurable": {"thread_id": request.thread_id}}
    try:
        graph.update_state(
            config,
            {
                "architecture_draft": result,
                "user_stories_draft": "",
            },
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to persist refined architecture: {exc}"),
        )

    _record_stage_event(request.thread_id, "architecture", "ai_revised")
    _record_stage_revision(
        request.thread_id,
        "architecture",
        "ai_revision",
        summary="Architecture updated from AI revision",
        instruction=request.instruction,
        reviewed=not request.preview_only,
        downstream_reset=STAGE_DOWNSTREAM["architecture"],
        content_length=len(result),
    )
    return RefineArchitectureResponse(architecture_draft=result)


# ---------------------------------------------------------------------------
# User Stories generation endpoint
# ---------------------------------------------------------------------------

class GenerateUserStoriesRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"


class GenerateUserStoriesResponse(BaseModel):
    user_stories_draft: str


@app.post("/api/generate_user_stories", response_model=GenerateUserStoriesResponse)
async def generate_user_stories(request: GenerateUserStoriesRequest):
    """
    Human-in-the-Loop gate: user approves the architecture, then this endpoint
    invokes the User Story Agent to produce epics and user stories with
    acceptance criteria and story points.
    """
    model_choice = request.model_choice.strip().lower()
    if model_choice not in VALID_MODEL_CHOICES:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_model_choice",
                f"Invalid model_choice '{model_choice}'. "
                f"Valid options: {sorted(VALID_MODEL_CHOICES)}",
            ),
        )

    config = {"configurable": {"thread_id": request.thread_id}}

    # Retrieve current state
    try:
        state_snapshot = graph.get_state(config)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to retrieve thread state: {exc}"),
        )

    if state_snapshot is None or not state_snapshot.values:
        raise HTTPException(
            status_code=400,
            detail=error_detail("missing_architecture", "Architecture is not ready yet."),
        )

    architecture_draft = state_snapshot.values.get("architecture_draft", "").strip()
    if not architecture_draft:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "missing_architecture",
                "Architecture must be generated before creating user stories.",
            ),
        )

    prd_draft = state_snapshot.values.get("prd_draft", "").strip()

    user_stories_prompt = build_user_stories_prompt(
        model_choice=model_choice,
        prd_draft=prd_draft,
        architecture_draft=architecture_draft,
    )

    # Invoke the model in a thread to avoid blocking the event loop
    try:
        result = await asyncio.to_thread(invoke_model, model_choice, user_stories_prompt)
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"User Story agent failed: {exc}"),
        )

    # Persist the user stories draft back into the graph state
    try:
        graph.update_state(config, {"user_stories_draft": result})
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to persist user stories draft: {exc}"),
        )

    _reset_stage_status(request.thread_id, "stories")
    _record_stage_event(request.thread_id, "stories", "generated")
    _record_stage_revision(
        request.thread_id,
        "stories",
        "generated",
        summary="User stories generated from architecture",
        reviewed=False,
        downstream_reset=STAGE_DOWNSTREAM["stories"],
        content_length=len(result),
    )
    return GenerateUserStoriesResponse(user_stories_draft=result)


class UpdateUserStoriesRequest(BaseModel):
    content: str
    change_source: str = "manual_edit"
    reviewed: bool = False
    instruction: str = ""
    change_context: str = ""


class UpdateUserStoriesResponse(BaseModel):
    success: bool
    user_stories_draft: str


@app.put("/api/user_stories/{thread_id}", response_model=UpdateUserStoriesResponse)
async def update_user_stories(thread_id: str, request: UpdateUserStoriesRequest):
    content = request.content.strip()
    if not content:
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "User stories content cannot be empty."),
        )
    change_source = _normalize_change_source(request.change_source)

    config = {"configurable": {"thread_id": thread_id}}
    try:
        graph.update_state(config, {"user_stories_draft": content})
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail(
                "state_error",
                f"Failed to persist user stories for thread '{thread_id}': {exc}",
            ),
        )
    _reset_stage_status(thread_id, "stories")
    _record_stage_event(thread_id, "stories", _event_type_for_change_source(change_source))
    _record_stage_revision(
        thread_id,
        "stories",
        change_source,
        summary=_default_revision_summary("stories", change_source, request.change_context),
        instruction=request.instruction,
        reviewed=request.reviewed,
        downstream_reset=STAGE_DOWNSTREAM["stories"],
        content_length=len(content),
    )
    return UpdateUserStoriesResponse(success=True, user_stories_draft=content)


class RefineUserStoriesRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    instruction: str
    preview_only: bool = False


class RefineUserStoriesResponse(BaseModel):
    user_stories_draft: str


@app.post("/api/refine_user_stories", response_model=RefineUserStoriesResponse)
async def refine_user_stories(request: RefineUserStoriesRequest):
    if not request.instruction.strip():
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "instruction cannot be empty."),
        )

    model_choice = validate_model_choice(request.model_choice)
    prd_draft = require_prd(request.thread_id)
    architecture_draft = require_architecture(request.thread_id)
    user_stories_draft = require_user_stories(request.thread_id)
    prompt = build_user_stories_refine_prompt(
        model_choice=model_choice,
        prd_draft=prd_draft,
        architecture_draft=architecture_draft,
        user_stories_draft=user_stories_draft,
        instruction=request.instruction.strip(),
    )

    try:
        result = await asyncio.to_thread(invoke_model, model_choice, prompt)
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"User stories refine failed: {exc}"),
        )

    if request.preview_only:
        return RefineUserStoriesResponse(user_stories_draft=result)

    config = {"configurable": {"thread_id": request.thread_id}}
    try:
        graph.update_state(config, {"user_stories_draft": result})
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to persist refined user stories: {exc}"),
        )

    _reset_stage_status(request.thread_id, "stories")
    _record_stage_event(request.thread_id, "stories", "ai_revised")
    _record_stage_revision(
        request.thread_id,
        "stories",
        "ai_revision",
        summary="User stories updated from AI revision",
        instruction=request.instruction,
        reviewed=not request.preview_only,
        downstream_reset=STAGE_DOWNSTREAM["stories"],
        content_length=len(result),
    )
    return RefineUserStoriesResponse(user_stories_draft=result)


# ---------------------------------------------------------------------------
# Stage chat endpoints
# ---------------------------------------------------------------------------

VALID_STAGES = {"architecture", "stories"}


class StageChatRequest(BaseModel):
    thread_id: str
    user_input: str
    model_choice: str = "ollama"
    preview_only: bool = False


class StageChatResponse(BaseModel):
    ai_response: str
    updated_content: Optional[str] = None


class StageHistoryMessage(BaseModel):
    role: str
    content: str


class StageHistoryResponse(BaseModel):
    messages: list[StageHistoryMessage]


@app.post("/api/stage/{stage}/chat", response_model=StageChatResponse)
async def stage_chat(stage: str, request: StageChatRequest):
    """
    Send a message to the stage-specific discussion agent.

    Each stage (architecture / stories) maintains its own conversation
    history, separate from the main PRD chat.  The agent has full context
    of the current PRD and relevant artifacts injected into every prompt.

    If the agent produces updated artifact content it wraps it in
    [CONTENT_START]…[CONTENT_END].  The endpoint parses this, persists the
    updated artifact back to the main thread state, and returns both the
    conversational message and the new content to the frontend.
    """
    if stage not in VALID_STAGES:
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_stage", f"Stage must be one of: {sorted(VALID_STAGES)}"),
        )
    if not request.user_input.strip():
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_input", "user_input cannot be empty."),
        )

    model_choice = validate_model_choice(request.model_choice)
    values = get_thread_values(request.thread_id)
    prd_draft           = str(values.get("prd_draft", "")).strip()
    architecture_draft  = str(values.get("architecture_draft", "")).strip()
    user_stories_draft  = str(values.get("user_stories_draft", "")).strip()

    # Build conversation text from persisted history
    history = _load_stage_messages(request.thread_id, stage)
    _append_stage_message(request.thread_id, stage, "user", request.user_input.strip())
    history.append({"role": "user", "content": request.user_input.strip()})

    conversation_text = "\n".join(
        f"{'User' if m['role'] == 'user' else 'Assistant'}: {m['content']}"
        for m in history
    )

    if stage == "architecture":
        prompt = build_arch_chat_prompt(
            model_choice=model_choice,
            prd_draft=prd_draft,
            architecture_draft=architecture_draft,
            conversation_text=conversation_text,
            latest_user_input=request.user_input.strip(),
        )
    else:
        prompt = build_stories_chat_prompt(
            model_choice=model_choice,
            prd_draft=prd_draft,
            architecture_draft=architecture_draft,
            user_stories_draft=user_stories_draft,
            conversation_text=conversation_text,
            latest_user_input=request.user_input.strip(),
        )

    try:
        raw_response = await asyncio.to_thread(invoke_model, model_choice, prompt)
    except (RuntimeError, ValueError) as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("model_error", f"Stage agent failed: {exc}"),
        )

    ai_message, updated_content = _extract_stage_content(raw_response)
    _append_stage_message(request.thread_id, stage, "assistant", ai_message)

    # Persist updated artifact back to main thread if the agent produced one
    # and the caller explicitly requested apply mode.
    if updated_content and not request.preview_only:
        main_config = {"configurable": {"thread_id": request.thread_id}}
        state_update = (
            {"architecture_draft": updated_content, "user_stories_draft": ""}
            if stage == "architecture"
            else {"user_stories_draft": updated_content}
        )
        try:
            graph.update_state(main_config, state_update)
            # Content was updated by AI — reset approval
            stage_key = "stories" if stage == "stories" else stage
            _reset_stage_status(request.thread_id, stage_key)
            if stage == "architecture":
                _reset_stage_status(request.thread_id, "stories")
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail=error_detail("state_error", f"Failed to persist updated {stage}: {exc}"),
            )

    if updated_content is not None and not request.preview_only:
        _record_stage_event(request.thread_id, stage, "ai_revised")
        _record_stage_revision(
            request.thread_id,
            stage,
            "ai_revision",
            summary=f"{stage.title()} updated from stage discussion",
            instruction=request.user_input.strip(),
            reviewed=False,
            downstream_reset=STAGE_DOWNSTREAM["stories"] if stage == "architecture" else STAGE_DOWNSTREAM[stage],
            content_length=len(updated_content),
        )
    return StageChatResponse(ai_response=ai_message, updated_content=updated_content)


@app.get("/api/stage/{stage}/chat/{thread_id}", response_model=StageHistoryResponse)
async def get_stage_chat_history(stage: str, thread_id: str):
    """Return the persisted stage chat history for a given thread."""
    if stage not in VALID_STAGES:
        raise HTTPException(
            status_code=400,
            detail=error_detail("invalid_stage", f"Stage must be one of: {sorted(VALID_STAGES)}"),
        )
    messages = _load_stage_messages(thread_id, stage)
    return StageHistoryResponse(
        messages=[StageHistoryMessage(role=m["role"], content=m["content"]) for m in messages]
    )


# ---------------------------------------------------------------------------
# Jira push endpoint
# ---------------------------------------------------------------------------

class PushToJiraRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    jira_domain: str
    jira_email: str
    jira_token: str
    jira_project_key: str


class PushToJiraResponse(BaseModel):
    success: bool
    created_issues: list[str]
    count: int


class PushToGitHubRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    github_owner: str
    github_repo: str
    github_token: str


class PushToGitHubResponse(BaseModel):
    success: bool
    created_issues: list[str]
    count: int


class DeliveryPreviewRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    target: str
    jira_project_key: str = ""
    github_owner: str = ""
    github_repo: str = ""


class DeliveryPreviewResponse(BaseModel):
    target: str
    items: list[dict]
    payload_preview: list[dict]


class DeliveryPublishRequest(BaseModel):
    thread_id: str
    model_choice: str = "ollama"
    target: str
    jira_domain: str = ""
    jira_email: str = ""
    jira_token: str = ""
    jira_project_key: str = ""
    github_owner: str = ""
    github_repo: str = ""
    github_token: str = ""


class DeliveryPublishResponse(BaseModel):
    success: bool
    target: str
    created_items: list[str]
    count: int


@app.post("/api/delivery/preview", response_model=DeliveryPreviewResponse)
async def preview_delivery(request: DeliveryPreviewRequest):
    model_choice = validate_model_choice(request.model_choice)
    target = validate_delivery_target(request.target)
    parsed_items = parse_delivery_items_or_raise(request.thread_id, model_choice)
    items = delivery_items_to_json(parsed_items)
    preview = DELIVERY_INTEGRATIONS[target].preview(
        parsed_items,
        build_delivery_config(
            target,
            require_credentials=False,
            jira_project_key=request.jira_project_key,
            github_owner=request.github_owner,
            github_repo=request.github_repo,
        ),
    )
    return DeliveryPreviewResponse(target=target, items=items, payload_preview=preview)


@app.post("/api/delivery/publish", response_model=DeliveryPublishResponse)
async def publish_delivery(request: DeliveryPublishRequest):
    model_choice = validate_model_choice(request.model_choice)
    target = validate_delivery_target(request.target)
    items = parse_delivery_items_or_raise(request.thread_id, model_choice)
    config = build_delivery_config(
        target,
        jira_domain=request.jira_domain,
        jira_email=request.jira_email,
        jira_token=request.jira_token,
        jira_project_key=request.jira_project_key,
        github_owner=request.github_owner,
        github_repo=request.github_repo,
        github_token=request.github_token,
    )
    try:
        result = DELIVERY_INTEGRATIONS[target].publish(items, config)
    except urllib.error.HTTPError as exc:
        error_body = exc.read().decode("utf-8", errors="replace")
        category = "jira_api_error" if target == "jira" else "github_api_error"
        raise HTTPException(
            status_code=502,
            detail=error_detail(category, f"{target} API error (HTTP {exc.code}): {error_body}"),
        ) from exc
    except urllib.error.URLError as exc:
        category = "jira_network_error" if target == "jira" else "github_network_error"
        raise HTTPException(
            status_code=502,
            detail=error_detail(category, f"Could not reach {target}: {exc.reason}"),
        ) from exc
    except Exception as exc:
        category = f"{target}_error"
        raise HTTPException(
            status_code=502,
            detail=error_detail(category, str(exc)),
        ) from exc

    return DeliveryPublishResponse(
        success=result.success,
        target=result.target,
        created_items=result.created,
        count=result.count,
    )


@app.post("/api/push_to_jira", response_model=PushToJiraResponse)
async def push_to_jira(request: PushToJiraRequest):
    result = await publish_delivery(
        DeliveryPublishRequest(
            thread_id=request.thread_id,
            model_choice=request.model_choice,
            target="jira",
            jira_domain=request.jira_domain,
            jira_email=request.jira_email,
            jira_token=request.jira_token,
            jira_project_key=request.jira_project_key,
        )
    )
    return PushToJiraResponse(
        success=result.success,
        created_issues=result.created_items,
        count=result.count,
    )


@app.post("/api/push_to_github", response_model=PushToGitHubResponse)
async def push_to_github(request: PushToGitHubRequest):
    result = await publish_delivery(
        DeliveryPublishRequest(
            thread_id=request.thread_id,
            model_choice=request.model_choice,
            target="github",
            github_owner=request.github_owner,
            github_repo=request.github_repo,
            github_token=request.github_token,
        )
    )
    return PushToGitHubResponse(
        success=result.success,
        created_issues=result.created_items,
        count=result.count,
    )


# ---------------------------------------------------------------------------
# Jira projects list endpoint
# ---------------------------------------------------------------------------

class JiraProject(BaseModel):
    key: str
    name: str
    id: str


class JiraProjectsResponse(BaseModel):
    projects: list[JiraProject]


@app.get("/api/jira/projects", response_model=JiraProjectsResponse)
async def get_jira_projects(domain: str, email: str, token: str):
    """
    Fetch all accessible Jira projects for the given credentials.
    Returns HTTP 401 on auth failure, HTTP 502 on network/other errors.
    """
    try:
        projects = list_jira_projects(domain, email, token)
        return JiraProjectsResponse(
            projects=[JiraProject(**project) for project in projects]
        )
    except urllib.error.HTTPError as exc:
        if exc.code in (401, 403):
            raise HTTPException(
                status_code=401,
                detail=error_detail("jira_auth_error", "Invalid Jira credentials"),
            )
        raise HTTPException(
            status_code=502,
            detail=error_detail("jira_api_error", f"Jira error: {exc.reason}"),
        )
    except urllib.error.URLError as exc:
        raise HTTPException(
            status_code=502,
            detail=error_detail(
                "jira_network_error",
                f"Could not reach Jira at '{domain}': {exc.reason}",
            ),
        )
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=error_detail("jira_error", str(exc)),
        )


# ---------------------------------------------------------------------------
# GitHub repos list endpoint
# ---------------------------------------------------------------------------

class GitHubRepo(BaseModel):
    full_name: str
    owner: str
    name: str


class GitHubReposResponse(BaseModel):
    repos: list[GitHubRepo]


@app.get("/api/github/repos", response_model=GitHubReposResponse)
async def get_github_repos(token: str):
    """Fetch all repos accessible by the given GitHub token."""
    try:
        repos = list_github_repos(token)
        return GitHubReposResponse(repos=[GitHubRepo(**r) for r in repos])
    except urllib.error.HTTPError as exc:
        if exc.code in (401, 403):
            raise HTTPException(
                status_code=401,
                detail=error_detail("github_auth_error", "Invalid GitHub token"),
            )
        raise HTTPException(
            status_code=502,
            detail=error_detail("github_api_error", f"GitHub error: {exc.reason}"),
        )
    except urllib.error.URLError as exc:
        raise HTTPException(
            status_code=502,
            detail=error_detail("github_network_error", f"Could not reach GitHub: {exc.reason}"),
        )
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=error_detail("github_error", str(exc)),
        )


# ---------------------------------------------------------------------------
# Stage status endpoints
# ---------------------------------------------------------------------------

VALID_STAGE_STATUSES = {"draft", "approved", "needs_revision"}


class StageStatusesResponse(BaseModel):
    prd: str
    architecture: str
    stories: str


class SetStageStatusRequest(BaseModel):
    status: str


@app.get("/api/stage/statuses/{thread_id}", response_model=StageStatusesResponse)
async def get_stage_statuses(thread_id: str):
    return StageStatusesResponse(
        prd=_get_stage_status(thread_id, "prd"),
        architecture=_get_stage_status(thread_id, "architecture"),
        stories=_get_stage_status(thread_id, "stories"),
    )


@app.patch("/api/stage/{stage}/status/{thread_id}")
async def set_stage_status_endpoint(stage: str, thread_id: str, request: SetStageStatusRequest):
    if stage not in ("prd", "architecture", "stories"):
        raise HTTPException(status_code=400, detail=error_detail("invalid_stage", "Invalid stage"))
    if request.status not in VALID_STAGE_STATUSES:
        raise HTTPException(status_code=400, detail=error_detail("invalid_status", f"Status must be one of: {sorted(VALID_STAGE_STATUSES)}"))
    _set_stage_status(thread_id, stage, request.status)
    event_label = "approved" if request.status == "approved" else "reopened" if request.status == "draft" else "marked_needs_revision"
    _record_stage_event(thread_id, stage, event_label)
    return {"stage": stage, "thread_id": thread_id, "status": request.status}


class StageEventItem(BaseModel):
    event_type: str
    detail: str
    created_at: float


class StageEventsResponse(BaseModel):
    events: list[StageEventItem]


class StageRevisionItem(BaseModel):
    id: int
    source: str
    summary: str
    instruction: str
    reviewed: bool
    downstream_reset: list[str]
    content_length: int
    created_at: float


class StageRevisionsResponse(BaseModel):
    revisions: list[StageRevisionItem]


class StageCommentItem(BaseModel):
    id: int
    body: str
    status: str
    created_at: float
    resolved_at: Optional[float] = None


class StageCommentsResponse(BaseModel):
    comments: list[StageCommentItem]


class CreateStageCommentRequest(BaseModel):
    body: str


class CreateStageCommentResponse(BaseModel):
    comment: StageCommentItem


class UpdateStageCommentRequest(BaseModel):
    status: str


class StageSummaryItem(BaseModel):
    stage: str
    status: str
    has_content: bool
    blocked_by: list[str]
    downstream_stages: list[str]
    downstream_impacted: list[str]
    stale: bool
    open_comments: int
    last_updated_at: Optional[float] = None
    last_revision_source: Optional[str] = None
    last_revision_summary: Optional[str] = None
    last_revision_reviewed: bool = False


class StageSummariesResponse(BaseModel):
    prd: StageSummaryItem
    architecture: StageSummaryItem
    stories: StageSummaryItem


@app.get("/api/stage/{stage}/events/{thread_id}", response_model=StageEventsResponse)
async def get_stage_events(stage: str, thread_id: str):
    if stage not in STAGE_KEYS:
        raise HTTPException(status_code=400, detail=error_detail("invalid_stage", "Invalid stage"))
    events = _load_stage_events(thread_id, stage)
    return StageEventsResponse(events=[StageEventItem(**e) for e in events])


@app.get("/api/stage/{stage}/revisions/{thread_id}", response_model=StageRevisionsResponse)
async def get_stage_revisions(stage: str, thread_id: str):
    if stage not in STAGE_KEYS:
        raise HTTPException(status_code=400, detail=error_detail("invalid_stage", "Invalid stage"))
    revisions = _load_stage_revisions(thread_id, stage)
    return StageRevisionsResponse(revisions=[StageRevisionItem(**revision) for revision in revisions])


@app.get("/api/stage/{stage}/comments/{thread_id}", response_model=StageCommentsResponse)
async def get_stage_comments(stage: str, thread_id: str):
    if stage not in STAGE_KEYS:
        raise HTTPException(status_code=400, detail=error_detail("invalid_stage", "Invalid stage"))
    comments = _load_stage_comments(thread_id, stage)
    return StageCommentsResponse(comments=[StageCommentItem(**comment) for comment in comments])


@app.post("/api/stage/{stage}/comments/{thread_id}", response_model=CreateStageCommentResponse)
async def create_stage_comment_endpoint(stage: str, thread_id: str, request: CreateStageCommentRequest):
    if stage not in STAGE_KEYS:
        raise HTTPException(status_code=400, detail=error_detail("invalid_stage", "Invalid stage"))
    body = request.body.strip()
    if not body:
        raise HTTPException(status_code=400, detail=error_detail("invalid_input", "Comment body cannot be empty."))
    comment = _create_stage_comment(thread_id, stage, body)
    _record_stage_event(thread_id, stage, "comment_added", body[:120])
    return CreateStageCommentResponse(comment=StageCommentItem(**comment))


@app.patch("/api/stage/comment/{comment_id}", response_model=CreateStageCommentResponse)
async def update_stage_comment_endpoint(comment_id: int, request: UpdateStageCommentRequest):
    status = request.status.strip().lower()
    if status not in {"open", "resolved"}:
        raise HTTPException(status_code=400, detail=error_detail("invalid_status", "Comment status must be 'open' or 'resolved'."))
    comment = _update_stage_comment_status(comment_id, status)
    if not comment:
        raise HTTPException(status_code=404, detail=error_detail("not_found", "Comment not found."))
    _record_stage_event(
        comment["thread_id"],
        comment["stage"],
        "comment_resolved" if status == "resolved" else "comment_reopened",
        comment["body"][:120],
    )
    return CreateStageCommentResponse(
        comment=StageCommentItem(
            id=comment["id"],
            body=comment["body"],
            status=comment["status"],
            created_at=comment["created_at"],
            resolved_at=comment["resolved_at"],
        )
    )


@app.get("/api/stage/summaries/{thread_id}", response_model=StageSummariesResponse)
async def get_stage_summaries(thread_id: str):
    artifacts = get_project_artifacts(thread_id)

    def build_summary(stage: str) -> StageSummaryItem:
        latest_revision = _get_latest_stage_revision(thread_id, stage)
        latest_timestamp = _get_latest_stage_timestamp(thread_id, stage)
        dependency_timestamps = [
            _get_latest_stage_timestamp(thread_id, dependency)
            for dependency in STAGE_DEPENDENCIES[stage]
        ]
        dependency_timestamps = [ts for ts in dependency_timestamps if ts is not None]
        stale = bool(
            _stage_has_content(artifacts, stage)
            and latest_timestamp is not None
            and any(ts > latest_timestamp for ts in dependency_timestamps)
        )
        downstream_impacted = [
            downstream
            for downstream in STAGE_DOWNSTREAM[stage]
            if _stage_has_content(artifacts, downstream)
            and latest_timestamp is not None
            and (
                (_get_latest_stage_timestamp(thread_id, downstream) or 0) < latest_timestamp
            )
        ]
        open_comments = sum(
            1 for comment in _load_stage_comments(thread_id, stage) if comment["status"] == "open"
        )
        return StageSummaryItem(
            stage=stage,
            status=_get_stage_status(thread_id, stage),
            has_content=_stage_has_content(artifacts, stage),
            blocked_by=[
                dependency for dependency in STAGE_DEPENDENCIES[stage]
                if not _stage_has_content(artifacts, dependency)
            ],
            downstream_stages=STAGE_DOWNSTREAM[stage],
            downstream_impacted=downstream_impacted,
            stale=stale,
            open_comments=open_comments,
            last_updated_at=latest_timestamp,
            last_revision_source=latest_revision["source"] if latest_revision else None,
            last_revision_summary=latest_revision["summary"] if latest_revision else None,
            last_revision_reviewed=latest_revision["reviewed"] if latest_revision else False,
        )

    return StageSummariesResponse(
        prd=build_summary("prd"),
        architecture=build_summary("architecture"),
        stories=build_summary("stories"),
    )


# ---------------------------------------------------------------------------
# PRD reset endpoint
# ---------------------------------------------------------------------------

class ResetPrdResponse(BaseModel):
    success: bool


@app.post("/api/reset_prd/{thread_id}", response_model=ResetPrdResponse)
async def reset_prd(thread_id: str):
    """
    Reset PRD and architecture state so the user can re-enter the clarification phase.
    Conversation history is preserved.
    """
    config = {"configurable": {"thread_id": thread_id}}
    try:
        graph.update_state(
            config,
            {
                "prd_draft": "",
                "is_ready_for_architecture": False,
                "architecture_draft": "",
                "user_stories_draft": "",
            },
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=error_detail("state_error", f"Failed to reset PRD for thread '{thread_id}': {exc}"),
        )
    return ResetPrdResponse(success=True)


# ---------------------------------------------------------------------------
# Export endpoint
# ---------------------------------------------------------------------------

@app.get("/api/export/{thread_id}")
async def export_project(thread_id: str, format: str = "markdown"):
    """
    Export the full project document (PRD + Architecture) as a downloadable Markdown file.
    Returns HTTP 400 if both sections are empty.
    """
    artifacts = get_project_artifacts(thread_id)

    if not artifacts.prd and not artifacts.architecture and not artifacts.user_stories:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "export_empty",
                "Nothing to export — PRD, architecture, and user stories are all empty.",
            ),
        )

    if format not in {"markdown", "json"}:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "invalid_export_format",
                f"Unsupported export format '{format}'. Valid options: ['json', 'markdown']",
            ),
        )

    if format == "json":
        return JSONResponse(content=export_project_json(artifacts))

    return Response(
        content=export_project_markdown(artifacts),
        media_type="text/markdown",
        headers={
            "Content-Disposition": f'attachment; filename="{thread_id}-project.md"'
        },
    )


# ---------------------------------------------------------------------------
# Model availability check
# ---------------------------------------------------------------------------

class ModelsCheckResponse(BaseModel):
    available: list[str]
    budgets: dict[str, dict[str, int]]


@app.get("/api/models/check", response_model=ModelsCheckResponse)
async def check_models():
    """
    Probe each supported AI backend and return which ones are reachable.
    Checks run concurrently to minimise latency.
    """
    checks = await asyncio.gather(
        *(asyncio.to_thread(adapter.is_available) for adapter in MODEL_ADAPTERS.values())
    )
    available = [
        adapter.model_choice
        for adapter, is_available in zip(MODEL_ADAPTERS.values(), checks)
        if is_available
    ]
    budgets = {
        adapter.model_choice: {
            "max_context_tokens": adapter.max_context_tokens,
            "prompt_budget_tokens": adapter.prompt_budget_tokens,
            "response_budget_tokens": adapter.response_budget_tokens,
        }
        for adapter in MODEL_ADAPTERS.values()
    }

    return ModelsCheckResponse(available=available, budgets=budgets)


@app.get("/health")
async def health_check():
    return {
        "status": "ok",
        "version": "2.0.0",
        "supported_models": sorted(VALID_MODEL_CHOICES),
    }


# ---------------------------------------------------------------------------
# File upload helpers
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".xlsx", ".xls", ".docx", ".pdf", ".md"}


def _extract_excel(data: bytes) -> str:
    """Convert all sheets of an Excel workbook to plain-text tables."""
    wb = openpyxl.load_workbook(filename=io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """Extract paragraphs and table cells from a Word document."""
    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        parts.append("")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))

    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    """Extract text from every page of a PDF using PyMuPDF."""
    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                parts.append(f"## Page {page_num}\n{text}")
    return "\n\n".join(parts)


def _extract_markdown(data: bytes) -> str:
    """Read a Markdown file as UTF-8 text."""
    return data.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# /api/upload endpoint
# ---------------------------------------------------------------------------

class UploadResponse(BaseModel):
    filename: str
    content: str


@app.post("/api/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename or "unknown"
    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=error_detail(
                "unsupported_file_type",
                f"Unsupported file type '{ext}'. "
                f"Supported: {sorted(SUPPORTED_EXTENSIONS)}",
            ),
        )

    try:
        raw = await file.read()
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=error_detail("file_read_error", f"Failed to read uploaded file: {exc}"),
        )

    try:
        if ext in {".xlsx", ".xls"}:
            content = _extract_excel(raw)
        elif ext == ".docx":
            content = _extract_docx(raw)
        elif ext == ".pdf":
            content = _extract_pdf(raw)
        elif ext == ".md":
            content = _extract_markdown(raw)
        else:
            # Should never reach here due to the extension check above
            raise HTTPException(status_code=415, detail="Unsupported file type.")
    except HTTPException:
        raise
    except Exception as exc:
            raise HTTPException(
                status_code=422,
                detail=error_detail("file_parse_error", f"Failed to parse '{filename}': {exc}"),
        )

    if not content.strip():
        raise HTTPException(
            status_code=422,
            detail=error_detail(
                "file_parse_error",
                f"No text content could be extracted from '{filename}'.",
            ),
        )

    return UploadResponse(filename=filename, content=content)
